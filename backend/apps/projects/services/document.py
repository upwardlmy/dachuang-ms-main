import io
from docx import Document  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
from apps.projects.models import Project

class DocumentService:
    @staticmethod
    def generate_project_doc(project_id):
        """
        Generate a Word document for the project application.
        """
        try:
            project = Project.objects.filter(id=project_id).select_related('leader', 'level', 'category', 'source').first()
            if not project:
                raise ValueError("Project not found")

            doc = Document()
            
            # Title
            title = doc.add_heading('大学生创新创业训练计划项目申报书', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Basic Info
            doc.add_heading('一、基本信息', level=1)
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            
            # Helper to set cell content
            def set_row(idx, label, value):
                row = table.rows[idx]
                row.cells[0].text = label
                row.cells[1].text = str(value) if value else ""

            set_row(0, "项目名称", project.title)
            set_row(1, "项目编号", project.project_no)
            set_row(2, "负责人", project.leader.real_name)
            set_row(3, "学号", project.leader.employee_id)
            set_row(4, "学院", getattr(project.leader, 'college', ''))
            set_row(5, "项目级别", project.level.label if project.level else "")

            # Project Content
            doc.add_heading('二、立项依据', level=1)
            doc.add_paragraph("（包含研究意义、现状分析等）")
            doc.add_paragraph(project.description)

            doc.add_heading('三、预期成果', level=1)
            doc.add_paragraph(project.expected_results)
            
            doc.add_heading('四、经费预算', level=1)
            doc.add_paragraph(f"申请经费：{project.budget} 元")
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer, f"{project.title}_申报书.docx"

        except Exception as e:
            print(f"Error generating doc: {e}")
            raise

    @staticmethod
    def generate_midterm_doc(project_id):
        return DocumentService._generate_generic_report(project_id, "中期检查报告")

    @staticmethod
    def generate_closure_doc(project_id):
        return DocumentService._generate_generic_report(project_id, "结题报告")

    @staticmethod
    def _generate_generic_report(project_id, report_type):
        try:
            project = Project.objects.get(id=project_id)
            doc = Document()
            doc.add_heading(f'{project.title} - {report_type}', 0)
            doc.add_paragraph(f"项目编号: {project.project_no}")
            doc.add_paragraph(f"负责人: {project.leader.real_name}")
            
            doc.add_heading('报告内容', level=1)
            doc.add_paragraph("此处为系统自动生成的报告摘要，详细内容请参考附件。")
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer, f"{project.title}_{report_type}.docx"
        except Exception as e:
            raise e
